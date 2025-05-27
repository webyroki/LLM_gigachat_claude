# MCP-сервер для генерации .docx файлов по шаблону
from context7.mcp.server.fastmcp import FastMCP
from docxtpl import DocxTemplate, TemplateError
import os
from datetime import datetime
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Инициализация MCP
mcp = FastMCP("DocGenerator")

def _validate_template_path(path: str) -> bool:
    """Проверяет существование и формат шаблона"""
    if not os.path.exists(path):
        raise FileNotFoundError(f"Шаблон не найден: {path}")
    if not path.lower().endswith('.docx'):
        raise ValueError("Шаблон должен быть в формате .docx")
    return True

@mcp.tool()
def generate_docx(path_template: str, context: dict, output_path: str) -> str:
    """
    Генерирует .docx-документ по шаблону с подстановкой переменных.
    Каждый файл получает уникальное имя с временной меткой (YYYYMMDD_HHMMSS).
    
    Args:
        path_template: Путь к файлу шаблона .docx
        context: Словарь с переменными для подстановки
        output_path: Базовый путь для сохранения файла
    
    Returns:
        Сообщение о результате операции
    """
    try:
        # Валидация шаблона
        _validate_template_path(path_template)
        
        # Создание уникального имени файла
        base, ext = os.path.splitext(output_path)
        if not ext:
            ext = '.docx'
        timestamp = datetime.now().strftime("_%Y%m%d_%H%M%S")
        final_path = f"{base}{timestamp}{ext}"
        
        # Создание документа
        doc = DocxTemplate(path_template)
        doc.render(context)
        
        # Создание директории если не существует
        output_dir = os.path.dirname(final_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        
        doc.save(final_path)
        
        logger.info(f"Документ успешно создан: {final_path}")
        return f"✅ Документ успешно создан: {final_path}"
        
    except FileNotFoundError as e:
        logger.error(f"Файл не найден: {e}")
        return f"❌ Файл не найден: {e}"
    except TemplateError as e:
        logger.error(f"Ошибка шаблона: {e}")
        return f"❌ Ошибка в шаблоне: {e}"
    except Exception as e:
        logger.error(f"Неожиданная ошибка: {e}")
        return f"❌ Ошибка генерации: {e}"

@mcp.tool()
def get_template_variables(path_template: str) -> str:
    """
    Возвращает список переменных, используемых в шаблоне .docx.
    
    Args:
        path_template: Путь к файлу шаблона .docx
    
    Returns:
        JSON-строка со списком переменных или сообщение об ошибке
    """
    try:
        _validate_template_path(path_template)
        
        doc = DocxTemplate(path_template)
        variables = list(doc.get_undeclared_template_variables())
        
        if not variables:
            return "ℹ️ В шаблоне не найдено переменных для подстановки"
        
        variables_str = ", ".join(f"'{var}'" for var in sorted(variables))
        return f"�� Переменные в шаблоне: {variables_str}"
        
    except Exception as e:
        logger.error(f"Ошибка получения переменных: {e}")
        return f"❌ Ошибка: {e}"

@mcp.tool()
def validate_template(path_template: str) -> str:
    """
    Проверяет корректность шаблона .docx.
    
    Args:
        path_template: Путь к файлу шаблона .docx
    
    Returns:
        Результат валидации шаблона
    """
    try:
        _validate_template_path(path_template)
        
        # Попытка загрузить шаблон
        doc = DocxTemplate(path_template)
        variables = list(doc.get_undeclared_template_variables())
        
        return f"✅ Шаблон корректен. Найдено переменных: {len(variables)}"
        
    except Exception as e:
        logger.error(f"Шаблон некорректен: {e}")
        return f"❌ Шаблон некорректен: {e}"

# === ДОПОЛНЕНИЕ: функции для работы с содержимым DOCX ===

@mcp.tool()
def read_docx(path: str) -> str:
    """
    Прочитать содержимое DOCX файла (дополнение).
    Args:
        path: Путь к файлу .docx
    Returns:
        Текстовое содержимое файла
    """
    try:
        from docx import Document
        if not os.path.exists(path):
            return f"❌ Файл '{path}' не найден."
        doc = Document(path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n".join(text) if text else "Документ пуст"
    except Exception as e:
        logger.error(f"Ошибка при чтении DOCX: {e}")
        return f"❌ Ошибка при чтении DOCX: {e}"

@mcp.tool()
def append_to_docx(path: str, text: str) -> str:
    """
    Добавить содержимое к существующему DOCX файлу (дополнение).
    Args:
        path: Путь к файлу .docx
        text: Текст для добавления
    Returns:
        Сообщение о результате
    """
    try:
        from docx import Document
        if not os.path.exists(path):
            return f"❌ Файл '{path}' не найден."
        doc = Document(path)
        doc.add_paragraph(text)
        doc.save(path)
        return f"✅ Текст успешно добавлен в '{path}'."
    except Exception as e:
        logger.error(f"Ошибка при добавлении текста в DOCX: {e}")
        return f"❌ Ошибка при добавлении текста: {e}"

# Точка входа MCP-сервера
if __name__ == "__main__":
    logger.info("Запуск MCP-сервера DocGenerator...")
    mcp.run(transport="stdio")