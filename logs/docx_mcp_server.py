#!/usr/bin/env python3
"""
MCP Server для создания и редактирования DOCX файлов
Требует: pip install python-docx mcp
"""

import asyncio
import json
import logging
import shutil
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches
from mcp.server import Server
from mcp.server.models import InitializationOptions
from mcp.server.stdio import stdio_server
from mcp.types import (
    Resource,
    Tool,
    TextContent,
    ImageContent,
    EmbeddedResource,
)

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("docx-mcp-server")

class DocxMCPServer:
    def __init__(self):
        self.server = Server("docx-server")
        self.setup_tools()
        self.setup_resources()

    def setup_tools(self):
        """Регистрация инструментов для работы с DOCX"""
        
        @self.server.list_tools()
        async def handle_list_tools() -> List[Tool]:
            return [
                Tool(
                    name="create_docx",
                    description="Создать новый DOCX документ с текстом и форматированием",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filename": {
                                "type": "string",
                                "description": "Имя файла (с расширением .docx)"
                            },
                            "title": {
                                "type": "string",
                                "description": "Заголовок документа"
                            },
                            "content": {
                                "type": "array",
                                "description": "Содержимое документа",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "type": {
                                            "type": "string",
                                            "enum": ["heading", "paragraph", "bullet_list", "numbered_list"]
                                        },
                                        "text": {"type": "string"},
                                        "level": {"type": "integer", "minimum": 1, "maximum": 3},
                                        "items": {
                                            "type": "array",
                                            "items": {"type": "string"}
                                        }
                                    }
                                }
                            }
                        },
                        "required": ["filename", "content"]
                    }
                ),
                Tool(
                    name="read_docx",
                    description="Прочитать содержимое DOCX файла",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filename": {
                                "type": "string",
                                "description": "Путь к DOCX файлу"
                            }
                        },
                        "required": ["filename"]
                    }
                ),
                Tool(
                    name="append_to_docx",
                    description="Добавить содержимое к существующему DOCX файлу",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filename": {
                                "type": "string",
                                "description": "Путь к существующему DOCX файлу"
                            },
                            "content": {
                                "type": "array",
                                "description": "Содержимое для добавления",
                                "items": {
                                    "type": "object",
                                    "properties": {
                                        "type": {
                                            "type": "string",
                                            "enum": ["heading", "paragraph", "bullet_list", "numbered_list"]
                                        },
                                        "text": {"type": "string"},
                                        "level": {"type": "integer", "minimum": 1, "maximum": 3},
                                        "items": {
                                            "type": "array",
                                            "items": {"type": "string"}
                                        }
                                    }
                                }
                            }
                        },
                        "required": ["filename", "content"]
                    }
                ),
                Tool(
                    name="list_docx_files",
                    description="Список всех DOCX файлов в указанной директории",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "directory": {
                                "type": "string",
                                "description": "Путь к директории (по умолчанию текущая)"
                            }
                        }
                    }
                ),
                Tool(
                    name="copy_docx",
                    description="Копировать DOCX файл в другую папку",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "source": {
                                "type": "string",
                                "description": "Путь к исходному DOCX файлу"
                            },
                            "destination": {
                                "type": "string",
                                "description": "Путь назначения (папка или полный путь к файлу)"
                            },
                            "new_name": {
                                "type": "string",
                                "description": "Новое имя файла (опционально)"
                            }
                        },
                        "required": ["source", "destination"]
                    }
                ),
                Tool(
                    name="move_docx",
                    description="Переместить DOCX файл в другую папку",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "source": {
                                "type": "string",
                                "description": "Путь к исходному DOCX файлу"
                            },
                            "destination": {
                                "type": "string",
                                "description": "Путь назначения (папка или полный путь к файлу)"
                            },
                            "new_name": {
                                "type": "string",
                                "description": "Новое имя файла (опционально)"
                            }
                        },
                        "required": ["source", "destination"]
                    }
                ),
                Tool(
                    name="create_folder",
                    description="Создать новую папку для DOCX файлов",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "folder_path": {
                                "type": "string",
                                "description": "Путь к новой папке"
                            }
                        },
                        "required": ["folder_path"]
                    }
                ),
                Tool(
                    name="delete_docx",
                    description="Удалить DOCX файл (перемещение в корзину если возможно)",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filename": {
                                "type": "string",
                                "description": "Путь к DOCX файлу для удаления"
                            },
                            "force": {
                                "type": "boolean",
                                "description": "Принудительное удаление без корзины",
                                "default": False
                            }
                        },
                        "required": ["filename"]
                    }
                )
            ]

        @self.server.call_tool()
        async def handle_call_tool(name: str, arguments: Dict[str, Any]) -> List[TextContent]:
            try:
                if name == "create_docx":
                    return await self.create_docx(**arguments)
                elif name == "read_docx":
                    return await self.read_docx(**arguments)
                elif name == "append_to_docx":
                    return await self.append_to_docx(**arguments)
                elif name == "list_docx_files":
                    return await self.list_docx_files(**arguments)
                elif name == "copy_docx":
                    return await self.copy_docx(**arguments)
                elif name == "move_docx":
                    return await self.move_docx(**arguments)
                elif name == "create_folder":
                    return await self.create_folder(**arguments)
                elif name == "delete_docx":
                    return await self.delete_docx(**arguments)
                else:
                    raise ValueError(f"Неизвестный инструмент: {name}")
            except Exception as e:
                logger.error(f"Ошибка при выполнении {name}: {e}")
                return [TextContent(type="text", text=f"Ошибка: {str(e)}")]

    def setup_resources(self):
        """Настройка ресурсов для доступа к файлам"""
        
        @self.server.list_resources()
        async def handle_list_resources() -> List[Resource]:
            return [
                Resource(
                    uri="file://docx",
                    name="DOCX Files",
                    description="Доступ к DOCX файлам в текущей директории",
                    mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            ]

    async def create_docx(self, filename: str, content: List[Dict], title: Optional[str] = None) -> List[TextContent]:
        """Создание нового DOCX документа"""
        try:
            doc = Document()
            
            # Добавляем заголовок документа если указан
            if title:
                doc.add_heading(title, 0)
            
            # Обрабатываем содержимое
            for item in content:
                item_type = item.get("type", "paragraph")
                text = item.get("text", "")
                level = item.get("level", 1)
                items = item.get("items", [])
                
                if item_type == "heading":
                    doc.add_heading(text, level)
                elif item_type == "paragraph":
                    doc.add_paragraph(text)
                elif item_type == "bullet_list":
                    for list_item in items:
                        doc.add_paragraph(list_item, style='List Bullet')
                elif item_type == "numbered_list":
                    for list_item in items:
                        doc.add_paragraph(list_item, style='List Number')
            
            # Сохраняем файл
            doc.save(filename)
            
            return [TextContent(
                type="text", 
                text=f"✅ DOCX файл '{filename}' успешно создан"
            )]
            
        except Exception as e:
            return [TextContent(
                type="text", 
                text=f"❌ Ошибка при создании файла: {str(e)}"
            )]

    async def read_docx(self, filename: str) -> List[TextContent]:
        """Чтение содержимого DOCX файла"""
        try:
            if not Path(filename).exists():
                return [TextContent(
                    type="text", 
                    text=f"❌ Файл '{filename}' не найден"
                )]
            
            doc = Document(filename)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Определяем тип параграфа
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name[-1] if paragraph.style.name[-1].isdigit() else "1"
                        content_parts.append(f"# Заголовок уровня {level}: {paragraph.text}")
                    elif paragraph.style.name == 'List Bullet':
                        content_parts.append(f"• {paragraph.text}")
                    elif paragraph.style.name == 'List Number':
                        content_parts.append(f"1. {paragraph.text}")
                    else:
                        content_parts.append(paragraph.text)
            
            content = "\n\n".join(content_parts) if content_parts else "Документ пуст"
            
            return [TextContent(
                type="text", 
                text=f"📄 Содержимое файла '{filename}':\n\n{content}"
            )]
            
        except Exception as e:
            return [TextContent(
                type="text", 
                text=f"❌ Ошибка при чтении файла: {str(e)}"
            )]

    async def append_to_docx(self, filename: str, content: List[Dict]) -> List[TextContent]:
        """Добавление содержимого к существующему DOCX файлу"""
        try:
            if not Path(filename).exists():
                return [TextContent(
                    type="text", 
                    text=f"❌ Файл '{filename}' не найден"
                )]
            
            doc = Document(filename)
            
            # Добавляем новое содержимое
            for item in content:
                item_type = item.get("type", "paragraph")
                text = item.get("text", "")
                level = item.get("level", 1)
                items = item.get("items", [])
                
                if item_type == "heading":
                    doc.add_heading(text, level)
                elif item_type == "paragraph":
                    doc.add_paragraph(text)
                elif item_type == "bullet_list":
                    for list_item in items:
                        doc.add_paragraph(list_item, style='List Bullet')
                elif item_type == "numbered_list":
                    for list_item in items:
                        doc.add_paragraph(list_item, style='List Number')
            
            # Сохраняем обновленный файл
            doc.save(filename)
            
            return [TextContent(
                type="text", 
                text=f"✅ Содержимое успешно добавлено в файл '{filename}'"
            )]
            
        except Exception as e:
            return [TextContent(
                type="text", 
                text=f"❌ Ошибка при добавлении содержимого: {str(e)}"
            )]

    async def list_docx_files(self, directory: str = ".") -> List[TextContent]:
        """Список DOCX файлов в директории"""
        try:
            path = Path(directory)
            if not path.exists():
                return [TextContent(
                    type="text", 
                    text=f"❌ Директория '{directory}' не найдена"
                )]
            
            docx_files = list(path.glob("*.docx"))
            
            if not docx_files:
                return [TextContent(
                    type="text", 
                    text=f"📁 В директории '{directory}' нет DOCX файлов"
                )]
            
            file_list = []
            for file in docx_files:
                size = file.stat().st_size
                size_str = f"{size} байт" if size < 1024 else f"{size/1024:.1f} КБ"
                file_list.append(f"📄 {file.name} ({size_str})")
            
            return [TextContent(
                type="text", 
                text=f"📁 DOCX файлы в '{directory}':\n\n" + "\n".join(file_list)
            )]
            
        except Exception as e:
            return [TextContent(
                type="text", 
                text=f"❌ Ошибка при получении списка файлов: {str(e)}"
            )]

    async def run(self):
        """Запуск сервера"""
        async with stdio_server() as (read_stream, write_stream):
            await self.server.run(
                read_stream,
                write_stream,
                InitializationOptions(
                    server_name="docx-mcp-server",
                    server_version="1.0.0",
                    capabilities=self.server.get_capabilities(
                        notification_options=None,
                        experimental_capabilities={}
                    )
                )
            )

async def main()